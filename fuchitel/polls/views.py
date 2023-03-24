from django.shortcuts import render, redirect, get_object_or_404
from django.core.mail import send_mail, EmailMessage
from django.conf import settings
from django.http import HttpResponse
from django.views.generic import ListView
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx import Document
from docx.shared import Inches

import os

from .models import Applicant

from .forms import ApplicantForm, EmailForm, SearchForm
from email.mime import application

def index(request):
    return render(request, 'polls/index.html')

def help_me(request):
    return render(request, 'polls/help.html')


def edit(request, pk):
    applicant = get_object_or_404(Applicant, pk=pk)
    if request.method == "POST":
        form = ApplicantForm(request.POST, request.FILES, instance=applicant)
        if form.is_valid():
            applicant = form.save(commit=False)
            print("img", request.FILES['image'])
            applicant.image = request.FILES['image']
            print(applicant.image)
            applicant.save()
        return redirect('detail', pk=applicant.pk)
    else:
        form = ApplicantForm(instance=applicant)
        context = {'applicant': applicant, 'form': form}
        return render(request, "polls/create.html", context=context)

def create(request):
    if request.method == "POST":
        form = ApplicantForm(request.POST)
        if form.is_valid():
            applicant = form.save()
            return redirect('detail', pk=applicant.pk)
    else:
        form = ApplicantForm()
        return render(request, "polls/create.html", {"form": form})

def detail(request, pk):
    form = EmailForm(request.POST)
    applicant = get_object_or_404(Applicant, pk=pk)
    context = {'applicant': applicant, 'form': form}
    return render(request, 'polls/detail.html', context=context)

def delete(request, pk):
    applicant = get_object_or_404(Applicant, pk=pk)
    applicant.delete()
    return redirect('index')

def get_doc(pk):
    applicant = get_object_or_404(Applicant, pk=pk)
    
    fio = (f"{applicant.family} " if applicant.family else "")
    fio = fio + f"{applicant.name} " if applicant.name else ""
    fio = fio + f"{applicant.patr}" if applicant.patr else ""
    
    import os
    logo = os.path.join(settings.MEDIA_ROOT, 'img', 'logo.png') 
    
    document = Document()
    
    p = document.add_paragraph()
    run = p.add_run()
    run.add_picture(logo, width=Inches(1.75))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    records = applicant.get_doc_model().items()

    table = document.add_table(rows=0, cols=2)
    
    # ФИО и фото
    row_cells = table.add_row().cells
    if applicant.image:
        print(applicant.image.path)
        p = row_cells[0].add_paragraph()
        run = p.add_run()
        run.add_picture(applicant.image.path, width=Inches(1.25))
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    else:
        row_cells[0].text = ''
    row_cells[1].text = fio
    row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # ДР если есть и возраст
    if applicant.bdate:
        row_cells = table.add_row().cells
        row_cells[0].text = 'Дата рождения'
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].text = f'{applicant.bdate.strftime("%d.%m.%Y")}  (Возраст: {applicant.age})'
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
    for key, value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value
        row_cells[0].paragraphs[0].runs[0].bold = True
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

 
    filename = os.path.join(settings.EMAIL_FILES, 'form.docx') 
    document.save(filename)

    return filename

def send(request, pk):
    applicant = get_object_or_404(Applicant, pk=pk)
    form = EmailForm(request.POST)
    print(form.is_valid())
    if form.is_valid():
        print("Валидна")
        emails = request.POST.get('email')
        print(emails)
        fio = applicant.fio
        email = EmailMessage(
            subject=f'Анкета {fio}',
            body='',
            from_email=settings.EMAIL_HOST_USER,
            to=[emails, settings.EMAIL_HOST_USER],
        )
        doc = get_doc(pk)
        email.attach_file(doc)
        email.send(fail_silently=True)
        
    return redirect('detail', pk=applicant.pk)

class SearchResultsView(ListView):
    model = Applicant
    template_name = 'polls/search.html'
    
    def get_queryset(self):
        query = self.request.GET.get('q')
        fio = query.split(' ')
        print(len(fio))
        qs = Applicant.objects.all()
        if len(fio) >= 1:
            if fio[0].isdigit():
                return qs.filter(questionnaire_num=fio[0])
            qs = qs.filter(family=fio[0].capitalize()) 
        if len(fio) >= 2:
            qs = qs.filter(name=fio[1].capitalize())
        if len(fio) >= 3:
            qs = qs.filter(patr=fio[2].capitalize())           
        return qs